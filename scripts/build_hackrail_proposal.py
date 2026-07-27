from __future__ import annotations

import shutil
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path('/Users/wolkebidian/Documents/AI-rail-guide')
REFERENCE = Path('/Users/wolkebidian/Downloads/HackRail 鐵客松產品或服務構想書.docx')
OUT_DIR = ROOT / 'outputs'
OUT = OUT_DIR / '軌語_RailTalk_HackRail產品服務構想書.docx'
ASSET_DIR = ROOT / 'outputs' / 'hackrail_proposal' / 'assets'
FONT_FILE = '/Library/Fonts/Arial Unicode.ttf'
FONT_NAME = 'Noto Sans TC'

NAVY = '12324A'
BLUE = '147D92'
TEAL = '20A39E'
PALE = 'EAF5F5'
PALE_BLUE = 'EAF1F7'
INK = '24333D'
GRAY = '667680'
WHITE = 'FFFFFF'
LINE = 'A9C4CC'


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_table_borders(table, color=LINE, size='6') -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = f'w:{edge}'
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), size)
        node.set(qn('w:color'), color)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement('w:cantSplit')
    tr_pr.append(cant)


def set_font(run, size=12, bold=False, color=INK, italic=False) -> None:
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(p, space_after=4, line=1.2, keep=False) -> None:
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep


def add_text(p, text, size=12, bold=False, color=INK, italic=False):
    r = p.add_run(text)
    set_font(r, size, bold, color, italic)
    return r


def add_hyperlink(p, text: str, url: str, size=10.5) -> None:
    part = p.part
    rid = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), rid)
    new_run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), BLUE); rpr.append(color)
    underline = OxmlElement('w:u'); underline.set(qn('w:val'), 'single'); rpr.append(underline)
    fonts = OxmlElement('w:rFonts')
    for key in ('ascii', 'hAnsi', 'eastAsia'):
        fonts.set(qn(f'w:{key}'), FONT_NAME)
    rpr.append(fonts)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rpr.append(sz)
    new_run.append(rpr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    p._p.append(hyperlink)


def clear_body(doc: Document) -> None:
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def heading(doc, text, level=1, kicker=None):
    if kicker:
        p = doc.add_paragraph()
        style_paragraph(p, 2, 1.0, True)
        add_text(p, kicker.upper(), 9.5, True, TEAL)
    p = doc.add_paragraph()
    style_paragraph(p, 9 if level == 1 else 5, 1.05, True)
    add_text(p, text, 20 if level == 1 else 14, True, NAVY)
    return p


def body(doc, text, bold_lead=None, after=5):
    p = doc.add_paragraph()
    style_paragraph(p, after, 1.28)
    if bold_lead and text.startswith(bold_lead):
        add_text(p, bold_lead, 12, True, NAVY)
        add_text(p, text[len(bold_lead):], 12)
    else:
        add_text(p, text, 12)
    return p


def bullet(doc, text, color=INK, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    style_paragraph(p, after, 1.2)
    add_text(p, '● ', 9, True, TEAL)
    add_text(p, text, 12, False, color)
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    add_text(p, '', 1)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def style_table(table, header=True, font_size=11.2, widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    if widths:
        set_col_widths(table, widths)
    for ri, row in enumerate(table.rows):
        prevent_row_split(row)
        if ri == 0 and header:
            set_repeat_table_header(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if ri == 0 and header:
                set_cell_shading(cell, NAVY)
            for p in cell.paragraphs:
                style_paragraph(p, 2, 1.1)
                for run in p.runs:
                    set_font(run, font_size, ri == 0 and header, WHITE if ri == 0 and header else INK)


def info_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], PALE_BLUE)
        for r in cells[0].paragraphs[0].runs:
            set_font(r, 11.5, True, NAVY)
        for r in cells[1].paragraphs[0].runs:
            set_font(r, 11.5)
    style_table(table, header=False, font_size=11.5, widths=[5.2, 10.8])
    return table


def add_page_header(doc, page_no, label):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Cm(12.7); t.columns[1].width = Cm(3.3)
    left, right = t.rows[0].cells
    left.text = label
    right.text = f'RAILTALK  /  {page_no:02d}'
    set_cell_shading(left, PALE); set_cell_shading(right, NAVY)
    for r in left.paragraphs[0].runs: set_font(r, 9.5, True, TEAL)
    for r in right.paragraphs[0].runs: set_font(r, 9.5, True, WHITE)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for c in (left, right): set_cell_margins(c, 80, 100, 80, 100)


def add_picture(doc, path, width=Inches(6.25), caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    shape = p.add_run().add_picture(str(path), width=width)
    doc_pr = shape._inline.docPr
    doc_pr.set('descr', caption or Path(path).stem)
    doc_pr.set('title', Path(path).stem)
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(c, 5, 1.0)
        add_text(c, caption, 9.5, False, GRAY)


def font(size, bold=False):
    return ImageFont.truetype(FONT_FILE, size=size)


def rounded(draw, box, fill, outline=None, radius=22, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def center_text(draw, box, text, fnt, fill, spacing=6):
    x1, y1, x2, y2 = box
    bb = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=spacing, align='center')
    tw, th = bb[2] - bb[0], bb[3] - bb[1]
    draw.multiline_text(((x1+x2-tw)/2, (y1+y2-th)/2), text, font=fnt, fill=fill, spacing=spacing, align='center')


def diagram_journey(path):
    im = Image.new('RGB', (1600, 720), 'white'); d = ImageDraw.Draw(im)
    d.text((70, 45), '旅客使用旅程｜從上車到地方探索', font=font(42, True), fill='#12324A')
    items = [('01', '開始行程', '選擇路線／車次'), ('02', '動態定位', '辨識目前站與下一站'), ('03', '主動導覽', '依窗外與行進階段說故事'), ('04', '語音插話', '旅客隨時提問'), ('05', '切站續聊', '更新 context，不沿用舊站'), ('06', '下車探索', '取得景點與步行建議')]
    x = 55
    for i, (n, title, desc) in enumerate(items):
        w = 235; y1=180; y2=545
        rounded(d, (x,y1,x+w,y2), '#EAF5F5' if i%2==0 else '#EAF1F7', '#A9C4CC')
        d.ellipse((x+82, y1+25, x+153, y1+96), fill='#147D92')
        center_text(d, (x+82,y1+25,x+153,y1+96), n, font(26,True), 'white')
        center_text(d, (x+20,y1+125,x+w-20,y1+205), title, font(30,True), '#12324A')
        center_text(d, (x+18,y1+225,x+w-18,y2-25), desc, font(23), '#40545F')
        if i < len(items)-1:
            d.line((x+w+5, 360, x+w+32, 360), fill='#20A39E', width=8)
            d.polygon([(x+w+32,348),(x+w+49,360),(x+w+32,372)], fill='#20A39E')
        x += 255
    d.text((70, 630), '關鍵體驗：旅客不用操作畫面，位置一變，導覽語境就跟著變。', font=font(27,True), fill='#147D92')
    im.save(path, quality=92)


def diagram_context(path):
    im = Image.new('RGB', (1600, 820), 'white'); d = ImageDraw.Draw(im)
    d.text((70, 45), '動態 Context 資料流｜讓每次回答站在「現在」', font=font(42,True), fill='#12324A')
    boxes = [
        (70,180,390,350,'位置／行程事件','手動切站・未來 GPS／TDX','#EAF1F7'),
        (470,150,1130,390,'JourneyContextSnapshot','journeyId  ·  revision  ·  currentStationId\nnextStationId  ·  phase  ·  guideSegmentIndex','#EAF5F5'),
        (1210,180,1530,350,'Realtime Session','session.update → session.updated','#EAF1F7'),
        (470,500,780,700,'回應關聯','correlation ID\n＋ context revision','#FFF4E8'),
        (830,500,1140,700,'本地工具','故事・POI・導覽段落\n沿用最新 revision','#FFF4E8'),
    ]
    for x1,y1,x2,y2,title,desc,fill in boxes:
        rounded(d,(x1,y1,x2,y2),fill,'#A9C4CC')
        center_text(d,(x1+15,y1+15,x2-15,y1+75),title,font(28,True),'#12324A')
        center_text(d,(x1+18,y1+85,x2-18,y2-18),desc,font(22),'#40545F')
    for a,b in [((390,265),(470,265)),((1130,265),(1210,265)),((800,390),(625,500)),((800,390),(985,500))]:
        d.line((*a,*b),fill='#20A39E',width=8)
    d.text((75,760),'順序保證：revision 遞增 → session.update → 等待確認 → response.create',font=font(26,True),fill='#147D92')
    im.save(path, quality=92)


def diagram_architecture(path):
    im=Image.new('RGB',(1600,820),'white'); d=ImageDraw.Draw(im)
    d.text((70,45),'系統技術架構｜開放資料 × 行程狀態 × 即時語音',font=font(42,True),fill='#12324A')
    cols=[(65,190,380,680,'資料層',['臺鐵車站／時刻表','即時位置／到離站','觀光景點資料','編輯查證內容']),(480,150,1120,720,'RailTalk Context Engine',['站點與階段判定','revision／correlation 管理','工具註冊與內容選擇','事件 trace／情境重播']),(1220,190,1535,680,'體驗層',['OpenAI Realtime API','文字 transcript','雙向串流語音','CLI → 未來正式 App'])]
    for ci,(x1,y1,x2,y2,title,items) in enumerate(cols):
        rounded(d,(x1,y1,x2,y2),'#EAF5F5' if ci==1 else '#EAF1F7','#A9C4CC')
        center_text(d,(x1+20,y1+22,x2-20,y1+90),title,font(30,True),'#12324A')
        y=y1+125
        for item in items:
            rounded(d,(x1+30,y,x2-30,y+78),'white','#D1E1E5',14,2)
            center_text(d,(x1+35,y+5,x2-35,y+73),item,font(22),'#40545F')
            y+=100
    for y in (310,470,630):
        d.line((385,y,468,y),fill='#20A39E',width=8); d.polygon([(468,y-12),(488,y),(468,y+12)],fill='#20A39E')
        d.line((1125,y,1208,y),fill='#20A39E',width=8); d.polygon([(1208,y-12),(1228,y),(1208,y+12)],fill='#20A39E')
    im.save(path, quality=92)


def diagram_evidence(path):
    im=Image.new('RGB',(1600,900),'#F5F7F8'); d=ImageDraw.Draw(im)
    d.text((65,45),'平溪線示範與 CLI 操作證據',font=font(42,True),fill='#12324A')
    rounded(d,(60,135,1540,410),'white','#A9C4CC')
    stations=['瑞芳','猴硐','三貂嶺','大華','十分','望古','嶺腳','平溪','菁桐']
    xs=[105+i*172 for i in range(9)]
    d.line((xs[0],255,xs[-1],255),fill='#147D92',width=12)
    for i,(x,s) in enumerate(zip(xs,stations)):
        d.ellipse((x-22,233,x+22,277),fill='#20A39E' if i not in (1,4,7) else '#F2A65A',outline='white',width=5)
        bb=d.textbbox((0,0),s,font=font(20,True)); d.text((x-(bb[2]-bb[0])/2,300),s,font=font(20,True),fill='#24333D')
    d.text((90,365),'手動／scenario 切站，驗證快速移動、導覽中切站、工具呼叫後切站與取消回應。',font=font(23),fill='#40545F')
    rounded(d,(60,455,1080,840),'#10212D','#10212D')
    terminal=[
        ('> /station houtong','#D8E7EB'),
        ('client  session.update    rev=2  station=houtong','#60D9C8'),
        ('server  session.updated   rev=2  station=houtong','#8EC5FF'),
        ('> 現在在哪裡？','#D8E7EB'),
        ('client  response.create   rev=2  corr=8d2…','#60D9C8'),
        ('assistant  我們現在來到猴硐，下一站是三貂嶺。','#FFF2C2'),
        ('trace saved  ·  audio content redacted','#8FA6B2'),
    ]
    y=490
    for line,color in terminal:
        d.text((95,y),line,font=font(23),fill=color); y+=45
    cards=[('20','離線測試'),('5','內建情境'),('1','唯一 Context')]
    y=470
    for n,label in cards:
        rounded(d,(1130,y,1535,y+100),'white','#A9C4CC',18,3)
        d.text((1160,y+20),n,font=font(42,True),fill='#147D92')
        d.text((1250,y+32),label,font=font(22,True),fill='#12324A')
        y+=125
    im.save(path, quality=92)


def add_footer_identity(doc):
    for section in doc.sections:
        section.footer_distance = Cm(0.8)
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        # Preserve existing page-number field when present; add identity before it.
        if not p.text.strip():
            add_text(p, '軌語 RailTalk｜行旅智能　', 9, False, GRAY)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True); ASSET_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REFERENCE, OUT)
    for name, fn in [('journey.png',diagram_journey),('context.png',diagram_context),('architecture.png',diagram_architecture),('evidence.png',diagram_evidence)]:
        fn(ASSET_DIR/name)

    doc=Document(OUT)
    clear_body(doc)
    sec=doc.sections[0]
    sec.page_width=Cm(21); sec.page_height=Cm(29.7)
    sec.left_margin=Cm(1.7); sec.right_margin=Cm(1.7); sec.top_margin=Cm(1.2); sec.bottom_margin=Cm(1.2)
    sec.header_distance=Cm(0.7); sec.footer_distance=Cm(0.75)
    normal=doc.styles['Normal']; normal.font.name=FONT_NAME; normal.font.size=Pt(12)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'),FONT_NAME)

    # PAGE 1 — official basic fields
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; style_paragraph(p,2,1.0)
    add_text(p,'HackRail 鐵客松',20,True,NAVY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; style_paragraph(p,7,1.0)
    add_text(p,'產品或服務構想書｜成果應用組',13,True,TEAL)
    info_table(doc,[('參賽編號','HR-17015'),('產品或服務名稱','軌語 RailTalk'),('公司／團隊名稱','行旅智能'),('預計使用本產品服務之鐵道機構','國營臺灣鐵路股份有限公司'),('產品或服務網址','https://github.com/Wolke/AI-rail-guide'),('測試用帳密','不適用；需自行設定 OPENAI_API_KEY')])
    heading(doc,'欲解決痛點',2)
    for t in ['既有服務聚焦班次與票務，缺少移動途中的地方理解。','固定音檔不能追問，也不會隨目前站與下一站調整。','快速切站時，AI 易沿用舊站語境，造成內容錯置。']:
        bullet(doc,t,after=2)
    heading(doc,'產品或服務簡介（300 字內）',2)
    body(doc,'軌語 RailTalk 結合鐵路開放資料、位置感知與生成式 AI。系統依目前站、下一站與行進階段主動提供沿線故事；旅客可用語音插話，並以最新 context revision 避免切站後沿用舊資訊。首波以平溪線驗證，讓車程成為可互動的地方旅行。',after=0)

    # PAGE 2 — official open data field
    page_break(doc); add_page_header(doc,2,'官方欄位｜開放資料來源')
    heading(doc,'開放資料來源（請詳列）',2,'DATA')
    data=[
        ('國營臺灣鐵路股份有限公司','臺鐵車站基本資料集','站點識別、名稱與路線基礎','https://data.gov.tw/dataset/33425'),
        ('國營臺灣鐵路股份有限公司','鐵路時刻表','車次、停靠順序與行程規劃','https://data.gov.tw/dataset/6138'),
        ('交通部','臺鐵列車即時位置動態資料','推估列車目前區間與位置','https://data.gov.tw/dataset/161161'),
        ('交通部','臺鐵列車即時到離站資料','校正到站、離站與延誤狀態','https://data.gov.tw/dataset/161156'),
        ('交通部觀光署','景點－觀光資訊資料庫','沿線景點與旅遊資訊候選','https://data.gov.tw/dataset/7777'),
        ('新北市政府觀光旅遊局','新北市觀光旅遊景點（中文）','平溪線周邊景點與地方資訊','https://data.gov.tw/dataset/122908'),
    ]
    table=doc.add_table(rows=1,cols=4)
    for i,h in enumerate(['提供機關','資料集名稱','RailTalk 用途','官方連結']): table.rows[0].cells[i].text=h
    for org,name,use,url in data:
        cells=table.add_row().cells
        cells[0].text=org; cells[1].text=name; cells[2].text=use
        p=cells[3].paragraphs[0]; add_hyperlink(p,'data.gov.tw',url,10)
    style_table(table,True,10.7,[3.8,4.4,5.4,2.4])
    p=doc.add_paragraph(); style_paragraph(p,0,1.0)
    add_text(p,'使用主辦單位釋出資料集：',11,True,NAVY); add_text(p,'無使用釋出資料集。',11)

    # PAGE 3 — official work description overview
    page_break(doc); add_page_header(doc,3,'官方欄位｜作品或服務說明')
    heading(doc,'作品或服務說明',1,'OVERVIEW')
    overview=[('一、緣起與創作目的','讓車程成為可對話的文化路徑。'),('二、市場調查與定位','補上訂票 App 之外的途中體驗。'),('三、使用對象與旅程情境','服務自由行、親子、銀髮與國際旅客。'),('四、服務特色與創新性','以版本化 Context 確保切站語境正確。'),('五、產品功能與技術架構','開放資料、Realtime 語音與可重播 trace。'),('六、未來規劃與擴充','串接 GPS／TDX、App、多語與場域實測。')]
    table=doc.add_table(rows=0,cols=2)
    for title,desc in overview:
        cells=table.add_row().cells; cells[0].text=title; cells[1].text=desc; set_cell_shading(cells[0],PALE_BLUE)
    style_table(table,False,11.2,[5.2,10.8])
    heading(doc,'是否曾獲獎／接受補助',2)
    body(doc,'未曾獲獎，亦未接受交通部鐵道局補助。')
    heading(doc,'成果現況與誠實邊界',2)
    body(doc,'已完成：macOS Realtime CLI、語音 I/O、切站、context revision、trace、五種情境與 20 項離線測試。')
    body(doc,'待完成：正式 App、GPS／TDX、多語、文史查證及場域實測。')

    # PAGE 4
    page_break(doc); add_page_header(doc,4,'01｜緣起與創作目的')
    heading(doc,'移動途中，正是故事最需要出現的時刻',1,'WHY NOW')
    body(doc,'現有服務善於回答「幾點到」，卻少有工具解釋窗外地景。軌語以唯一 JourneyContextSnapshot 管理目前站、下一站、行進階段與導覽段落；每次變動產生新 revision，舊回應不再推進新行程。')
    add_picture(doc,ASSET_DIR/'journey.png',Inches(5.8),'圖 1｜位置一變，導覽語境就跟著變。')
    heading(doc,'創作目的',2)
    for t in ['提升沿線文化與地方景點的可發現性。','以雙向語音降低操作門檻，並避免把舊位置說成現在。']:
        bullet(doc,t,after=2)

    # PAGE 5
    page_break(doc); add_page_header(doc,5,'02｜市場調查與定位')
    heading(doc,'不做另一個訂票 App，而是補上「途中體驗層」',1,'POSITIONING')
    body(doc,'RailTalk 與既有服務是互補關係：臺鐵 e 訂通解決行前與行程管理；一般語音導覽平台提供景點內容；RailTalk 則以鐵路行程 context 串起移動中的主動導覽與雙向問答。')
    comp=[('服務','核心任務','鐵路行程 Context','雙向即時語音','開放資料整合'),('臺鐵 e 訂通','查詢、訂票、乘車資訊','有行程／票務資訊','非文化導覽主軸','以營運資訊為主'),('izi.TRAVEL','景點與城市語音導覽','非鐵路狀態核心','以既有導覽內容為主','依內容供應者'),('軌語 RailTalk','移動中的 AI 語音伴旅','目前站＋下一站＋階段＋revision','可插話、可追問、切站續聊','車站、即時列車、景點資料')]
    table=doc.add_table(rows=0,cols=5)
    for row in comp:
        cells=table.add_row().cells
        for i,v in enumerate(row): cells[i].text=v
    style_table(table,True,10.1,[3.0,3.9,3.4,3.0,3.2])
    heading(doc,'優先市場與價值主張',2)
    body(doc,'B2G／B2B 提供鐵道與觀光組織可策展的底座；旅客獲得免手持、可追問的陪伴；地方端把「經過」轉化為「願意下車」。')
    body(doc,'場域測試目標：站點正確率 ≥95%、首包延遲 <2.5 秒；並追蹤導覽完成率、追問率與旅客滿意度。')

    # PAGE 6
    page_break(doc); add_page_header(doc,6,'03｜使用對象與旅程情境')
    heading(doc,'同一段鐵路，因旅客與當下狀態而不同',1,'JOURNEY')
    personas=[('自由行','理解沿線、臨時下車','下一站亮點＋步行景點'),('親子','故事化、可追問','分段導覽＋問題引導'),('銀髮','少操作、到站提醒','免手持＋慢速摘要'),('國際旅客','語言與文化脈絡','未來多語＋地名對照')]
    table=doc.add_table(rows=1,cols=3)
    for i,h in enumerate(['使用對象','旅程需求','RailTalk 回應方式']): table.rows[0].cells[i].text=h
    for a,b,c in personas:
        cells=table.add_row().cells; cells[0].text=a; cells[1].text=b; cells[2].text=c
    style_table(table,True,11,[3.4,6.2,6.5])
    heading(doc,'平溪線示範旅程',2)
    steps=[('瑞芳出發','建立 journey 與支線背景。'),('猴硐進站','revision 更新，回答景點追問。'),('十分前後','依 approaching／at_station 切換。'),('平溪至菁桐','依時間提供下車探索。')]
    for a,b in steps:
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(.3); style_paragraph(p,5,1.22)
        add_text(p,a+'｜',12,True,TEAL); add_text(p,b,12)
    heading(doc,'服務原則',2)
    body(doc,'旅客問題優先；位置可信度不足時採保守描述並告知資料狀態。')

    # PAGE 7
    page_break(doc); add_page_header(doc,7,'04｜服務特色與創新性')
    heading(doc,'創新不只在語音，而在語音背後的狀態一致性',1,'INNOVATION')
    add_picture(doc,ASSET_DIR/'context.png',Inches(5.65),'圖 2｜每次回答與工具續答都帶著 context revision。')
    features=[('唯一真相','Snapshot 集中管理站點、階段與段落。'),('確認後回應','session.updated 後才建立新回應。'),('回應可追蹤','correlation ID 關聯取消、工具與完成。'),('可重現除錯','JSONL trace 與情境腳本重播失敗。')]
    table=doc.add_table(rows=0,cols=2)
    for title,desc in features:
        c=table.add_row().cells; c[0].text=title; c[1].text=desc; set_cell_shading(c[0],PALE)
    style_table(table,False,11,[3.6,12.5])
    body(doc,'差異在於處理追問、生成期間切站、工具續答與舊回應晚到，而非只做 GPS 觸發音檔。')

    # PAGE 8
    page_break(doc); add_page_header(doc,8,'05｜產品功能與技術架構')
    heading(doc,'從資料到聲音，每一層都能被觀察與測試',1,'SYSTEM')
    add_picture(doc,ASSET_DIR/'architecture.png',Inches(5.6),'圖 3｜目前以 CLI 驗證核心；正式 App 為下一階段。')
    heading(doc,'目前原型功能',2)
    for t in ['24 kHz mono PCM16 麥克風、Realtime PCM 喇叭播放與文字 transcript。','切站／階段命令、五種情境腳本與本地故事／POI 工具。','trace 遮蔽 API Key 與音訊；helper 失敗時降級純文字。']:
        bullet(doc,t,after=2)
    heading(doc,'技術選擇',2)
    body(doc,'Node.js 22+ 以 WebSocket 直連 Realtime API；Swift／CoreAudio 負責 macOS 音訊。20 項一般測試不耗額度，live smoke test 需明確啟用。')

    # PAGE 9
    page_break(doc); add_page_header(doc,9,'06｜未來規劃與擴充')
    heading(doc,'先證明語境可靠，再把體驗帶進真實列車',1,'ROADMAP')
    add_picture(doc,ASSET_DIR/'evidence.png',Inches(5.55),'圖 4｜平溪線九站示範與可重現 CLI 情境。')
    roadmap=[('現在｜已驗證','CLI、revision、語音 I/O、trace、五情境與 20 測試。'),('0–3 個月','GPS／TDX、資料可信度、文史查證。'),('3–6 個月','正式 App、多語、平溪線場域實測。'),('6–12 個月','路線擴展、內容管理、成效儀表板與 API。')]
    table=doc.add_table(rows=0,cols=2)
    for a,b in roadmap:
        c=table.add_row().cells; c[0].text=a; c[1].text=b; set_cell_shading(c[0],PALE_BLUE)
    style_table(table,False,10.9,[4.4,11.7])
    body(doc,'風險對策：定位不準採保守描述；網路中斷用下載摘要；生成內容須有來源與審核；語音與位置資料最小化。')
    p=doc.add_paragraph(); style_paragraph(p,0,1.0)
    add_text(p,'原型：',9.5,True,GRAY); add_hyperlink(p,'github.com/Wolke/AI-rail-guide','https://github.com/Wolke/AI-rail-guide',9.5)
    add_text(p,'　來源：政府資料開放平臺、',9.5,False,GRAY)
    add_hyperlink(p,'臺鐵 e 訂通','https://apps.apple.com/tw/app/%E5%8F%B0%E9%90%B5e%E8%A8%82%E9%80%9A/id1441617748',9.5)
    add_text(p,'、',9.5,False,GRAY); add_hyperlink(p,'izi.TRAVEL','https://www.izi.travel/en/app',9.5)

    add_footer_identity(doc)
    # Global table font and paragraph cleanup.
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.size is None: set_font(r,12)
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build()
